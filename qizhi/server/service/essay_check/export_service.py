"""
论文审查报告导出服务：使用 python-docx 生成 Word 文件，支持多文件打包为 zip。
"""

import io
import re
import zipfile
from typing import Any

import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from common.utils.logger import get_logger

logger = get_logger(__name__)

# ── 清理文本中的 XML 不兼容字符 ──
# XML 允许: \x09 (tab), \x0A (\n), \x0D (\r)，其余控制字符及 \x00 必须清除
_XML_BAD_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]")


def _clean_text(text: str) -> str:
    """移除 XML 不兼容的控制字符，防止 python-docx 报错。"""
    if not text:
        return ""
    return _XML_BAD_CHARS.sub("", str(text))

# ── 颜色常量 ──
COLOR_HEADER = RGBColor(0x33, 0x33, 0x33)
COLOR_PASS = RGBColor(0x67, 0xC2, 0x3A)
COLOR_FAIL = RGBColor(0xF5, 0x6C, 0x6C)
COLOR_GRAY = RGBColor(0x90, 0x93, 0x99)


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色。"""
    shading = docx.oxml.OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, font_name: str = "微软雅黑", size: int | None = None, color: RGBColor | None = None, bold: bool = False):
    """设置 run 的字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold


def _add_styled_paragraph(doc: docx.Document, text: str, size: int = 10, bold: bool = False, color: RGBColor | None = None, space_after: int = 6):
    """添加带样式的段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(_clean_text(text))
    _set_run_font(run, size=size, bold=bold, color=color)
    return p


def _add_heading_styled(doc: docx.Document, text: str, level: int = 1):
    """添加样式化的标题。"""
    heading = doc.add_heading(_clean_text(text), level=level)
    for run in heading.runs:
        _set_run_font(run, size=14 if level == 1 else 12, bold=True, color=COLOR_HEADER)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def _create_table(doc: docx.Document, headers: list[str], rows: list[list[str]], col_widths: list[Inches] | None = None):
    """创建检查点表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 表头
    header_fill_color = "DEE8FF"
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = _clean_text(header)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_run_font(run, size=10, bold=True, color=COLOR_HEADER)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
        _set_cell_shading(cell, header_fill_color)

    # 数据行
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = _clean_text(str(cell_text))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=9)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
            # 检查不通过的行加红色背景
            if col_idx == 1 and cell_text == "不通过":
                _set_cell_shading(cell, "FDECEA")

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    return table


def _filter_pages_for_export(pages: list, export_scope: str) -> list:
    """按导出范围过滤逐页数据。"""
    if export_scope == "summary_only":
        return []
    if not pages:
        return []
    if export_scope == "summary_problem_pages":
        return [
            p for p in pages
            if isinstance(p, dict) and (
                p.get("status") != "completed"
                or any(
                    not cp.get("passed")
                    for cp in (p.get("check_results") or [])
                    if isinstance(cp, dict)
                )
            )
        ]
    return pages


def generate_report_docx(report_data: dict, export_scope: str = "summary_only") -> bytes:
    """
    将单个论文审查报告数据生成 Word 文件字节内容。

    Args:
        report_data: 包含 filename, overall_score, summary, check_domains,
                     reference_issues, recommendations, pages 的字典。
        export_scope: summary_only | summary_problem_pages | full

    Returns:
        docx 文件的字节内容。
    """
    doc = docx.Document()

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(10)

    # ── 标题 ──
    filename = report_data.get("filename", "未知文件")
    title = doc.add_heading("论文分析报告", level=0)
    for run in title.runs:
        _set_run_font(run, size=18, bold=True, color=COLOR_HEADER)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_styled_paragraph(doc, f"文件名：{filename}", size=11, color=COLOR_GRAY, space_after=12)

    # ── 总体评分 + 摘要 ──
    _add_heading_styled(doc, "总体评价", level=1)

    score = _clean_text(str(report_data.get("overall_score", "N/A")))
    score_p = doc.add_paragraph()
    score_p.paragraph_format.space_after = Pt(6)
    score_label = score_p.add_run("总体评分：")
    _set_run_font(score_label, size=11, bold=True)
    score_value = score_p.add_run(score)
    score_color = COLOR_PASS if score == "合格" else (COLOR_FAIL if score == "不合格" else COLOR_FAIL)
    _set_run_font(score_value, size=14, bold=True, color=score_color)

    summary = report_data.get("summary", "")
    if summary:
        _add_styled_paragraph(doc, summary, size=10, space_after=12)

    # ── 检查域 ──
    check_domains = report_data.get("check_domains") or []
    if check_domains:
        _add_heading_styled(doc, "检查详情", level=1)
        for domain in check_domains:
            if not isinstance(domain, dict):
                continue
            domain_name = domain.get("domain", "")
            _add_heading_styled(doc, domain_name, level=2)

            check_points = domain.get("check_points", []) or []
            if check_points:
                headers = ["检查点", "结果", "详情"]
                rows = []
                for cp in check_points:
                    if not isinstance(cp, dict):
                        continue
                    rows.append([
                        cp.get("name", cp.get("check_point", "")),
                        "通过" if cp.get("passed") else "不通过",
                        cp.get("detail", "") or "-",
                    ])
                _create_table(
                    doc, headers, rows,
                    col_widths=[Inches(2), Inches(0.8), Inches(3.5)],
                )

    # ── 参考文献问题 ──
    ref_issues = report_data.get("reference_issues")
    if ref_issues:
        _add_heading_styled(doc, "参考文献问题", level=1)
        for issue in ref_issues:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(_clean_text(str(issue)))
            _set_run_font(run, size=10)

    # ── 修改建议 ──
    recommendations = report_data.get("recommendations")
    if recommendations:
        _add_heading_styled(doc, "修改建议", level=1)
        for rec in recommendations:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(_clean_text(str(rec)))
            _set_run_font(run, size=10)

    # ── 逐页检查结果 ──
    pages = _filter_pages_for_export(report_data.get("pages") or [], export_scope)

    if pages:
        doc.add_page_break()
        _add_heading_styled(doc, "逐页检查结果", level=1)

        page_type_map = {
            "cover": "封面", "commitment": "承诺书", "abstract_cn": "中文摘要",
            "abstract_en": "英文摘要", "toc": "目录", "body": "正文",
            "conclusion": "结论", "references": "参考文献", "appendix": "附录",
            "task_sheet": "任务书", "assessment_form": "考核表",
            "expert_review": "专家评阅", "defense_record": "答辩记录", "unknown": "未知",
        }

        for page in pages:
            if not isinstance(page, dict):
                continue
            page_num = page.get("page_number", "?")
            page_type = page_type_map.get(page.get("page_type", ""), page.get("page_type", "未知"))
            page_status = page.get("status", "")

            page_heading = doc.add_paragraph()
            page_heading.paragraph_format.space_before = Pt(8)
            page_heading.paragraph_format.space_after = Pt(4)
            run_page = page_heading.add_run(f"第 {page_num} 页")
            _set_run_font(run_page, size=11, bold=True)
            run_type = page_heading.add_run(f"  |  {page_type}  |  {page_status}")
            _set_run_font(run_type, size=10, color=COLOR_GRAY)

            check_results = page.get("check_results") or []
            if check_results:
                cp_headers = ["检查点", "结果", "详情"]
                cp_rows = []
                for cp in check_results:
                    if not isinstance(cp, dict):
                        continue
                    cp_rows.append([
                        cp.get("check_point", cp.get("name", "")),
                        "通过" if cp.get("passed") else "不通过",
                        cp.get("detail", "") or "-",
                    ])
                _create_table(
                    doc, cp_headers, cp_rows,
                    col_widths=[Inches(2), Inches(0.8), Inches(3.5)],
                )

    # ── 输出字节 ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_zip(task_reports: list[dict], export_scope: str = "summary_only") -> bytes:
    """
    将多个报告数据打包为 zip 文件。

    Args:
        task_reports: 列表，每项包含 {filename, report_data}。
        export_scope: summary_only | summary_problem_pages | full

    Returns:
        zip 文件的字节内容。
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in task_reports:
            filename = item["filename"]
            report_data = item["report_data"]
            # 生成安全的文件名
            safe_name = _sanitize_filename(filename)
            docx_bytes = generate_report_docx(report_data, export_scope=export_scope)
            zf.writestr(f"{safe_name}.docx", docx_bytes)
    return buf.getvalue()


def export_to_zip_from_bytes(docx_files: list[tuple[str, bytes]]) -> bytes:
    """
    将已生成的 docx 字节内容打包为 zip 文件。

    Args:
        docx_files: 列表，每项为 (文件名, docx字节内容)。

    Returns:
        zip 文件的字节内容。
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname, fbytes in docx_files:
            zf.writestr(fname, fbytes)
    return buf.getvalue()


def _sanitize_filename(filename: str) -> str:
    """清理文件名，移除特殊字符并截断过长名称。"""
    # 移除扩展名
    name = filename.rsplit(".", 1)[0] if "." in filename else filename
    # 移除不安全字符
    for ch in r'\/:*?"<>|':
        name = name.replace(ch, "_")
    # 截断到 80 字符
    if len(name) > 80:
        name = name[:77] + "..."
    return name
