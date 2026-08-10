import asyncio
from io import BytesIO
import sys
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from course_generation_workflow import (
    build_course_blueprint_from_plan,
    build_node_generation_context,
    normalize_course_plan_contract,
)
from course_quality import build_grounding_quality_report
from material_evidence import (
    attach_evidence_to_plan,
    build_evidence_catalog_summary,
    extract_grounding_annotations,
)
import material_pipeline
from material_pipeline import prepare_course_materials
import material_parser
from material_models import MaterialAsset
from material_parser import PdfPageOcrParser
from material_storage import MaterialRepository, MaterialStorageError


class FakeUpload:
    def __init__(self, filename: str, content: bytes, content_type: str = "text/markdown"):
        self.filename = filename
        self.content_type = content_type
        self._content = content
        self._offset = 0

    async def read(self, size: int) -> bytes:
        if self._offset >= len(self._content):
            return b""
        chunk = self._content[self._offset:self._offset + size]
        self._offset += len(chunk)
        return chunk


@pytest.mark.asyncio
async def test_multiple_materials_parse_with_bounded_parallelism_and_stable_order(tmp_path, monkeypatch):
    repository = MaterialRepository(tmp_path / "materials")
    assets = [
        await repository.create_text_asset(filename=f"source-{index}.md", content=f"# 资料 {index}\n\n正文 {index}")
        for index in range(4)
    ]
    original_parse = material_pipeline.parse_material_asset
    active = 0
    max_active = 0

    async def observed_parse(current_repository, asset):
        nonlocal active, max_active
        active += 1
        max_active = max(max_active, active)
        try:
            await asyncio.sleep(0.02)
            return await original_parse(current_repository, asset)
        finally:
            active -= 1

    monkeypatch.setenv("MATERIAL_PARSE_CONCURRENCY", "2")
    monkeypatch.setattr(material_pipeline, "parse_material_asset", observed_parse)

    prepared = await prepare_course_materials(
        course_id="course-parallel",
        material_bindings=[{"asset_id": asset.asset_id} for asset in assets],
        legacy_materials=[],
        repository=repository,
    )

    assert max_active == 2
    assert [item["asset_id"] for item in prepared["material_assets"]] == [asset.asset_id for asset in assets]
    assert all(item["quality_state"] == "ready" for item in prepared["parsed_documents"])


@pytest.mark.asyncio
async def test_upload_parse_evidence_and_cache(tmp_path):
    repository = MaterialRepository(tmp_path / "materials")
    content = "# 导数\n\n定义：导数刻画瞬时变化率。\n\n题目：根据定义求导？\n"
    first = await repository.save_upload(FakeUpload("calculus.md", content.encode()))
    second = await repository.save_upload(FakeUpload("same.md", content.encode()))

    assert first.asset_id == second.asset_id
    prepared = await prepare_course_materials(
        course_id="course-1",
        material_bindings=[{
            "asset_id": first.asset_id,
            "purpose": "content_source",
            "priority": "core",
            "authority": "primary",
            "usage_policy": "must_use",
        }],
        legacy_materials=[],
        repository=repository,
    )
    prepared_again = await prepare_course_materials(
        course_id="course-2",
        material_bindings=prepared["material_bindings"],
        legacy_materials=[],
        repository=repository,
    )

    assert prepared["parsed_documents"][0]["parse_status"] == "parsed"
    assert prepared["evidence_catalog"]
    assert prepared["evidence_catalog"][0]["source_text"]
    assert prepared["evidence_catalog"][0]["locator"]["section_path"] == ["导数"]
    assert prepared_again["parsed_documents"][0]["document_id"] == prepared["parsed_documents"][0]["document_id"]
    assert repository.get_asset(first.asset_id).bound_course_ids == ["course-1", "course-2"]


@pytest.mark.asyncio
async def test_docling_slim_parses_real_docx_into_evidence(tmp_path):
    stream = BytesIO()
    document = Document()
    document.add_heading("Derivative", level=1)
    document.add_paragraph("A derivative describes an instantaneous rate of change.")
    document.save(stream)

    repository = MaterialRepository(tmp_path / "materials")
    asset = await repository.save_upload(FakeUpload(
        "derivative.docx",
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ))
    prepared = await prepare_course_materials(
        course_id="course-docx",
        material_bindings=[{"asset_id": asset.asset_id}],
        legacy_materials=[],
        repository=repository,
    )

    parsed = prepared["parsed_documents"][0]
    assert parsed["parse_status"] == "parsed"
    assert parsed["parser_name"] == "docling"
    assert any("instantaneous rate of change" in item["source_text"] for item in prepared["evidence_catalog"])


@pytest.mark.asyncio
async def test_upload_rejects_path_traversal_and_fake_pdf(tmp_path):
    repository = MaterialRepository(tmp_path / "materials")
    with pytest.raises(MaterialStorageError, match="文件名不安全"):
        await repository.save_upload(FakeUpload("../bad.md", b"# bad"))
    with pytest.raises(MaterialStorageError, match="有效 PDF"):
        await repository.save_upload(FakeUpload("fake.pdf", b"not a pdf", "application/pdf"))


@pytest.mark.asyncio
async def test_evidence_is_assigned_by_node_instead_of_broadcast(tmp_path):
    repository = MaterialRepository(tmp_path / "materials")
    asset = await repository.save_upload(FakeUpload(
        "calculus.md",
        "# 导数\n\n定义：导数是瞬时变化率。\n\n# 积分\n\n定义：积分描述累积量。".encode(),
    ))
    prepared = await prepare_course_materials(
        course_id="course-grounded",
        material_bindings=[{
            "asset_id": asset.asset_id,
            "purpose": "content_source",
            "priority": "core",
            "authority": "primary",
            "usage_policy": "must_use",
        }],
        legacy_materials=[],
        repository=repository,
    )
    plan = normalize_course_plan_contract({
        "course_title": "微积分",
        "chapters": [{
            "title": "基础",
            "sections": [
                {"title": "导数定义", "learning_objective": "解释导数", "assessment": ["解释变化率"]},
                {"title": "概率导论", "learning_objective": "解释概率", "assessment": ["计算概率"]},
            ],
        }],
    })
    plan, coverage = attach_evidence_to_plan(
        plan,
        evidence=prepared["evidence_catalog"],
        bindings=prepared["material_bindings"],
    )
    first, second = plan["chapters"][0]["sections"]

    assert first["evidence_refs"]
    assert second["evidence_refs"] == []
    assert "material_refs" not in first
    assert coverage["asset_coverage"][0]["assigned_nodes"] == ["L2-1-1"]

    artifacts = {
        **prepared,
        "course_generation_brief": {"subject": "微积分"},
        "subject_pedagogy_profile": {},
        "difficulty_profile": {},
        "evidence_coverage_plan": coverage,
    }
    blueprint = build_course_blueprint_from_plan(plan, artifacts)
    context = build_node_generation_context(
        course_metadata={**artifacts, "course_blueprint": blueprint},
        node=blueprint["nodes"][0],
    )
    assert "当前节点限定证据包" in context
    assert prepared["evidence_catalog"][0]["evidence_id"] in context


def test_grounding_markers_are_extracted_and_reported():
    evidence_id = "ev-abc123"
    content, annotations, invalid = extract_grounding_annotations(
        f"导数描述瞬时变化率。[[evidence:{evidence_id}]]",
        {evidence_id},
    )
    assert "[[evidence:" not in content
    assert annotations[0]["evidence_id"] == evidence_id
    assert invalid == []

    course = {
        "material_assets": [{"asset_id": "mat-1", "filename": "导数.md", "status": "parsed"}],
        "material_bindings": [{"asset_id": "mat-1", "purpose": "content_source", "usage_policy": "must_use"}],
        "evidence_catalog": [{"evidence_id": evidence_id, "asset_id": "mat-1"}],
        "evidence_coverage_plan": {"asset_coverage": [{"asset_id": "mat-1", "assigned_nodes": ["L2-1-1"]}]},
        "nodes": [{
            "node_id": "L2-1-1",
            "node_level": 2,
            "grounding_contract": {"required_evidence_ids": [evidence_id]},
            "grounding_annotations": annotations,
        }],
    }
    report = build_grounding_quality_report(course)
    assert report["passed"] is True
    assert report["material_coverage"][0]["coverage_level"] == "used"


def test_material_upload_api_uses_persisted_asset(monkeypatch, tmp_path):
    from routers import materials

    repository = MaterialRepository(tmp_path / "materials")
    monkeypatch.setattr(materials, "material_repository", repository)
    app = FastAPI()
    app.include_router(materials.router, prefix="/api")
    client = TestClient(app)

    response = client.post(
        "/api/materials",
        files={"file": ("notes.md", b"# Notes\n\nEvidence.", "text/markdown")},
    )
    assert response.status_code == 201
    asset_id = response.json()["asset_id"]
    assert repository.get_asset(asset_id) is not None
    assert "source_name" not in response.json()

    parsed = client.post(f"/api/materials/{asset_id}/parse")
    assert parsed.status_code == 200
    assert parsed.json()["quality_report"]["schema_version"] == "parsed_document_quality_v2"
    assert parsed.json()["quality_report"]["status"] == "ready"
    assert parsed.json()["preview"][0]["locator"]["section_path"] == ["Notes"]

    delete = client.delete(f"/api/materials/{asset_id}")
    assert delete.status_code == 200


def test_scanned_pdf_fallback_preserves_page_locators_and_quality(tmp_path, monkeypatch):
    source = tmp_path / "scan.pdf"
    source.write_bytes(b"%PDF-1.4\n")
    asset = MaterialAsset(
        asset_id="mat-scan",
        filename="scan.pdf",
        extension=".pdf",
        mime_type="application/pdf",
        detected_mime="application/pdf",
        size_bytes=source.stat().st_size,
        sha256="scan-sha",
        source_name="source.pdf",
        uploaded_at="2026-08-10T00:00:00+00:00",
        updated_at="2026-08-10T00:00:00+00:00",
    )

    class FakePage:
        def render(self, *, scale):
            assert scale == 2.0
            return self

        def to_pil(self):
            return self

        def save(self, path, *, format):
            assert format == "PNG"
            path.write_bytes(b"png")

        def close(self):
            return None

    class FakePdf:
        def __init__(self, _path):
            self.pages = [FakePage(), FakePage()]

        def __len__(self):
            return len(self.pages)

        def __getitem__(self, index):
            return self.pages[index]

        def close(self):
            return None

    monkeypatch.setitem(sys.modules, "pypdfium2", SimpleNamespace(PdfDocument=FakePdf))
    monkeypatch.setattr(
        material_parser,
        "_ocr_image",
        lambda path: [{
            "text": f"第 {path.stem.split('-')[-1]} 页内容",
            "confidence": 0.92,
            "bbox": [0, 0, 100, 20],
        }],
    )

    document = PdfPageOcrParser().parse(asset, source)

    assert document.parse_status == "parsed"
    assert document.parser_name == "pdf_page_ocr"
    assert [block.locator.page for block in document.blocks] == [1, 2]
    assert document.quality["ocr_page_coverage"] == 1.0
    assert document.warnings


def test_build_evidence_catalog_summary_covers_every_asset_not_just_earliest():
    """Regression test for the outline-planning truncation bias bug.

    20 assets uploaded in order, each producing a different number of
    evidence items (early ones intentionally produce many more than later
    ones, mimicking "20 years of exam papers uploaded oldest-first"). With
    a naive `evidence[:max_items]` truncation, the first couple of assets
    alone would fill the entire max_items=80 budget and later assets would
    be completely invisible to outline planning. The fix must guarantee
    every asset gets at least one representative entry in the summary.
    """
    evidence: list[dict] = []
    asset_ids = [f"asset-{index:02d}" for index in range(20)]
    for position, asset_id in enumerate(asset_ids):
        # Earlier assets (smaller position) produce far more evidence blocks.
        count = 30 if position < 3 else 2
        for item_index in range(count):
            evidence.append({
                "evidence_id": f"ev-{asset_id}-{item_index}",
                "asset_id": asset_id,
                "kind": "claim",
                "summary": f"{asset_id} 证据 {item_index}",
                "locator": {},
            })

    summary = build_evidence_catalog_summary(evidence, max_items=80)

    for asset_id in asset_ids:
        assert asset_id in summary, f"{asset_id} should have at least one representative evidence entry"

    # Sanity: still respects the max_items cap on number of lines.
    assert len(summary.splitlines()) <= 80


def test_build_evidence_catalog_summary_prefers_higher_priority_within_asset():
    evidence = [
        {
            "evidence_id": "ev-low",
            "asset_id": "asset-x",
            "kind": "claim",
            "summary": "low priority",
            "priority": "supporting",
            "authority": "context_only",
            "locator": {},
        },
        {
            "evidence_id": "ev-high",
            "asset_id": "asset-x",
            "kind": "claim",
            "summary": "high priority",
            "priority": "core",
            "authority": "primary",
            "locator": {},
        },
    ]
    summary = build_evidence_catalog_summary(evidence, max_items=1)
    assert "ev-high" in summary
    assert "ev-low" not in summary


def test_build_evidence_catalog_summary_no_truncation_needed():
    evidence = [
        {"evidence_id": f"ev-{i}", "asset_id": "asset-a", "kind": "claim", "summary": "x", "locator": {}}
        for i in range(5)
    ]
    summary = build_evidence_catalog_summary(evidence, max_items=80)
    assert len(summary.splitlines()) == 5
