"""Deterministic teaching-calendar exports from one saved calendar revision."""

from __future__ import annotations

import csv
import math
import unicodedata
from copy import copy
from io import BytesIO, StringIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


COLUMNS = [
    ("记录", "sequence"),
    ("日期", "date"),
    ("教学内容", "content_summary"),
    ("教学要求（含作业）", "requirements"),
    ("上课地点", "location"),
    ("上课教师", "teacher_name"),
    ("教学类型", "teaching_type"),
    ("实验小组", "group_code"),
    ("教学时数", "credit_hours"),
    ("备注", "notes"),
]
DOCX_WIDTHS_MM = [9, 19, 38, 43, 21, 15, 15, 12, 10, 12]
XLSX_WIDTHS = [8, 14, 34, 40, 20, 15, 14, 12, 10, 18]


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _session_rows(calendar: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for record_number, session in enumerate(calendar.get("sessions") or [], start=1):
        row: list[str] = []
        for _, key in COLUMNS:
            if key == "sequence":
                row.append(str(record_number))
            elif key == "date":
                date_value = _text(session.get("date"))
                start = _text(session.get("start_time"))[:5]
                end = _text(session.get("end_time"))[:5]
                time_value = f"\n{start}—{end}" if start and end else ""
                row.append(f"{date_value}{time_value}")
            else:
                row.append(_text(session.get(key)))
        rows.append(row)
    return rows


def _metadata(calendar: dict[str, Any], course: dict[str, Any]) -> dict[str, str]:
    request = course.get("generation_request") or {}
    brief = request.get("teacher_course_brief") or course.get("teacher_course_brief") or {}
    teachers = sorted({str(item.get("teacher_name") or "").strip() for item in calendar.get("sessions") or [] if str(item.get("teacher_name") or "").strip()})
    return {
        "course_code": _text(brief.get("course_code") or course.get("course_code")),
        "course_name": _text(calendar.get("course_title") or course.get("course_name") or "未命名课程"),
        "credits": _text(brief.get("credits") or brief.get("credit") or ""),
        "weekly_hours": _text(brief.get("weekly_hours") or brief.get("weekly_class_hours") or ""),
        "teachers": "、".join(teachers),
        "offering_number": _text(brief.get("course_offering_number") or brief.get("selection_number") or ""),
        "academic_year": _text(calendar.get("academic_year")),
        "term": _text(calendar.get("term")),
    }


def _set_cell_text(cell, value: str, *, size: float = 6.5, bold: bool = False, centered: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def build_docx(calendar: dict[str, Any], course: dict[str, Any]) -> bytes:
    metadata = _metadata(calendar, course)
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(f"浙江大学本科教学日历（{metadata['academic_year']}{metadata['term']}学期）")
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(13)

    meta = document.add_table(rows=2, cols=8)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta_widths = [15, 31, 15, 46, 12, 20, 15, 40]
    for row in meta.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Mm(meta_widths[index])
    labels = [
        (meta.cell(0, 0), "课程代码", True), (meta.cell(0, 1), metadata["course_code"], False),
        (meta.cell(0, 2), "课程名称", True), (meta.cell(0, 3), metadata["course_name"], False),
        (meta.cell(0, 4), "学分", True), (meta.cell(0, 5), metadata["credits"], False),
        (meta.cell(0, 6), "周学时", True), (meta.cell(0, 7), metadata["weekly_hours"], False),
        (meta.cell(1, 0), "主讲教师", True), (meta.cell(1, 1).merge(meta.cell(1, 3)), metadata["teachers"], False),
        (meta.cell(1, 4), "选课课号", True), (meta.cell(1, 5).merge(meta.cell(1, 7)), metadata["offering_number"], False),
    ]
    for cell, value, bold in labels:
        _set_cell_text(cell, value, size=6.5, bold=bold, centered=True)

    document.add_paragraph().paragraph_format.space_after = Pt(1)
    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for index, (label, _) in enumerate(COLUMNS):
        table.columns[index].width = Mm(DOCX_WIDTHS_MM[index])
        table.cell(0, index).width = Mm(DOCX_WIDTHS_MM[index])
        _set_cell_text(table.cell(0, index), label, size=6.2, bold=True, centered=True)
    _set_repeat_header(table.rows[0])
    for values in _session_rows(calendar):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = Mm(DOCX_WIDTHS_MM[index])
            _set_cell_text(cells[index], value, size=6.2, centered=index in {0, 1, 5, 6, 7, 8})
    notes_row = table.add_row().cells
    notes_label = notes_row[0].merge(notes_row[1])
    notes_value = notes_row[2].merge(notes_row[-1])
    _set_cell_text(notes_label, "备注", size=6.2, bold=True, centered=True)
    _set_cell_text(notes_value, _text(calendar.get("notes")), size=6.2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(metadata["offering_number"])
    footer_run.font.name = "宋体"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    footer_run.font.size = Pt(6)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_xlsx(calendar: dict[str, Any], course: dict[str, Any]) -> bytes:
    metadata = _metadata(calendar, course)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "教学日历"
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=10)
    sheet.cell(1, 1, f"浙江大学本科教学日历（{metadata['academic_year']}{metadata['term']}学期）")
    sheet.cell(1, 1).font = Font(name="宋体", size=14, bold=True)
    sheet.cell(1, 1).alignment = Alignment(horizontal="center")
    sheet.cell(2, 1, "课程代码")
    sheet.cell(2, 2, metadata["course_code"])
    sheet.cell(2, 3, "课程名称")
    sheet.merge_cells("D2:E2")
    sheet.cell(2, 4, metadata["course_name"])
    sheet.cell(2, 6, "学分")
    sheet.cell(2, 7, metadata["credits"])
    sheet.cell(2, 8, "周学时")
    sheet.merge_cells("I2:J2")
    sheet.cell(2, 9, metadata["weekly_hours"])
    sheet.cell(3, 1, "主讲教师")
    sheet.merge_cells("B3:E3")
    sheet.cell(3, 2, metadata["teachers"])
    sheet.cell(3, 6, "选课课号")
    sheet.merge_cells("G3:J3")
    sheet.cell(3, 7, metadata["offering_number"])
    headers = [label for label, _ in COLUMNS]
    for column, value in enumerate(headers, start=1):
        cell = sheet.cell(4, column, value)
        cell.font = Font(name="宋体", size=10, bold=True)
        cell.fill = PatternFill("solid", fgColor="EDEDFC")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row_index, values in enumerate(_session_rows(calendar), start=5):
        for column, value in enumerate(values, start=1):
            sheet.cell(row_index, column, value).alignment = Alignment(vertical="top", wrap_text=True)
        # Excel does not reliably auto-size wrapped rows when opened outside
        # desktop Excel.  Estimate the two long text columns explicitly so the
        # last lines remain visible in browser previews and printed exports.
        wrapped_lines = max(
            _xlsx_wrapped_lines(values[2], XLSX_WIDTHS[2]),
            _xlsx_wrapped_lines(values[3], XLSX_WIDTHS[3]),
            _xlsx_wrapped_lines(values[1], XLSX_WIDTHS[1]),
        )
        sheet.row_dimensions[row_index].height = min(150, max(34, wrapped_lines * 13.5 + 8))
    thin = Side(style="thin", color="B8BECC")
    for row in sheet.iter_rows(min_row=2, max_row=max(4, 4 + len(calendar.get("sessions") or [])), min_col=1, max_col=10):
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            font = copy(cell.font)
            font.name = "宋体"
            font.size = 9
            cell.font = font
    for index, width in enumerate(XLSX_WIDTHS, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A5"
    sheet.print_title_rows = "4:4"
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _xlsx_wrapped_lines(value: str, column_width: float) -> int:
    """Estimate wrapped Excel lines using East Asian display width."""
    capacity = max(1, int(column_width))
    total = 0
    for source_line in (_text(value).splitlines() or [""]):
        display_width = sum(
            2 if unicodedata.east_asian_width(char) in {"W", "F", "A"} else 1
            for char in source_line
        )
        total += max(1, math.ceil(display_width / capacity))
    return total


def build_csv(calendar: dict[str, Any]) -> bytes:
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow([label for label, _ in COLUMNS])
    writer.writerows(_session_rows(calendar))
    return ("\ufeff" + output.getvalue()).encode("utf-8")


def build_pdf(calendar: dict[str, Any], course: dict[str, Any]) -> bytes:
    metadata = _metadata(calendar, course)
    output = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CalendarTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=13, leading=16, alignment=TA_CENTER, spaceAfter=6)
    cell_style = ParagraphStyle("CalendarCell", parent=styles["BodyText"], fontName="STSong-Light", fontSize=5.8, leading=7.2, alignment=TA_LEFT)
    center_style = ParagraphStyle("CalendarCenter", parent=cell_style, alignment=TA_CENTER)
    document = SimpleDocTemplate(output, pagesize=A4, leftMargin=8 * mm, rightMargin=8 * mm, topMargin=14 * mm, bottomMargin=14 * mm, title=f"{metadata['course_name']} 教学日历", author=metadata["teachers"])
    story: list[Any] = [Paragraph(f"浙江大学本科教学日历（{metadata['academic_year']}{metadata['term']}学期）", title_style)]
    meta_data = [
        ["课程代码", metadata["course_code"], "课程名称", metadata["course_name"], "学分", metadata["credits"], "周学时", metadata["weekly_hours"]],
        ["主讲教师", metadata["teachers"], "", "", "选课课号", metadata["offering_number"], "", ""],
    ]
    meta_table = Table([[Paragraph(_text(value), center_style) for value in row] for row in meta_data], colWidths=[15*mm,31*mm,15*mm,46*mm,12*mm,20*mm,15*mm,40*mm], hAlign="CENTER")
    meta_table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.35,colors.black),("SPAN",(1,1),(3,1)),("SPAN",(5,1),(7,1)),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BACKGROUND",(0,0),(-1,-1),colors.white)]))
    story.extend([meta_table, Spacer(1, 2 * mm)])
    data = [[Paragraph(label, center_style) for label, _ in COLUMNS]]
    for values in _session_rows(calendar):
        data.append([Paragraph(value.replace("\n", "<br/>"), center_style if index in {0,1,5,6,7,8} else cell_style) for index, value in enumerate(values)])
    data.append([Paragraph("备注", center_style), "", Paragraph(_text(calendar.get("notes")), cell_style), "", "", "", "", "", "", ""])
    table = Table(data, colWidths=[width * mm for width in DOCX_WIDTHS_MM], hAlign="CENTER", repeatRows=1)
    table_style = [("GRID",(0,0),(-1,-1),0.3,colors.black),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#F1F2F6")),("SPAN",(0,-1),(1,-1)),("SPAN",(2,-1),(-1,-1))]
    table.setStyle(TableStyle(table_style))
    story.append(table)

    def footer(canvas, _document):
        canvas.saveState()
        canvas.setFont("STSong-Light", 6)
        canvas.setFillColor(colors.HexColor("#8A91A3"))
        canvas.drawRightString(A4[0] - 8 * mm, 7 * mm, metadata["offering_number"])
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()


EXPORTERS = {"docx": build_docx, "xlsx": build_xlsx, "pdf": build_pdf}
