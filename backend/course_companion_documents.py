"""Template-backed formal documents that sit outside the teaching-content chain.

Companion documents are course-owned, versioned formal files.  A template owns
the input contract, deterministic document structure, validation and export
layout; generated DOCX files are projections and never become a second truth.
"""

from __future__ import annotations

import hashlib
import json
import re
from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from storage import DATA_DIR

COMPANION_DOCUMENT_SCHEMA = "course_companion_document_v1"
COMPANION_TEMPLATE_SCHEMA = "course_companion_template_v1"
_ID_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9._-]{0,199}")


GRADING_COMPONENTS = [
    {
        "component_id": "classroom_participation",
        "name": "课堂讨论、交流",
        "weight": 20,
        "scope": "个人",
        "details": "考勤、课堂交流、Workshop表现（10分）\n讨论1：破冰+初识设计思维（3分，口述）\n讨论2：身边的设计思维（3分，口述）\n讨论3：再谈设计思维（4分，PPT）",
    },
    {
        "component_id": "stage_group_work",
        "name": "实践成果—阶段评价",
        "weight": 16,
        "scope": "分组小作业",
        "details": "选题分享：突出有意义（4分，PPT）\n定题分享：突出有挑战（4分，PPT）\n开题分享：突出有价值（4分，PPT）\n激励项：约3次大班分享机会（3分，PPT）\n注：原样例上述分项合计15分，请任课教师确认剩余1分的给分方式。",
    },
    {
        "component_id": "final_group_project",
        "name": "实践成果—最终评价",
        "weight": 24,
        "scope": "分组大作业",
        "details": "教师评分（6分）\n助教评分（6分）\n组间评分（6分）\n组内评价（6分，结合助教、教师评价）\n提交物：15页PPT以内",
    },
    {
        "component_id": "online_learning",
        "name": "线上学习—中国大学MOOC",
        "weight": 20,
        "scope": "个人",
        "details": "以平台记录的课程学习与测验成绩为依据。",
    },
    {
        "component_id": "offline_exam",
        "name": "线下考试",
        "weight": 20,
        "scope": "个人",
        "details": "开卷考试。",
    },
]


CHECKLIST_ITEMS = [
    ("actual_papers", "实考试卷按实际份数上交存档"),
    ("paper_a", "空白试卷A卷、A卷标准答案（含参考答案、评分细则）各一份上交存档"),
    ("paper_b", "空白试卷B卷、B卷标准答案（含参考答案、评分细则）各一份上交存档；多个平行班由出卷老师提供"),
    ("grade_sheets", "含平时成绩、期末成绩和总评成绩的成绩单2份，签名后上交存档"),
    ("exam_analysis", "试卷分析填写规范完整，签名后上交存档"),
    ("teaching_calendar", "从教务系统导出本学期教学日历，打印上交存档"),
    ("grading_rubric", "课程成绩评定细则完整，与教学大纲一致，签名上交存档"),
    ("grade_breakdown", "课程成绩明细汇总表与成绩评定细则对应"),
    ("syllabus_alignment", "考核方式、成绩评定方式与课程教学大纲实际情况相符"),
    ("paper_duplication", "近三年A卷重复率不超过30%，当年A卷、B卷不重复，卷首课程信息完整"),
    ("marking_signature", "考生实考试卷卷首评阅人签全名，卷首和卷面标明各题得分"),
    ("total_signature", "考生实考试卷卷首登分表总分处由任课老师签全名"),
    ("room_records", "考场情况记录表、考场签到表上交存档"),
]


COMPANION_DOCUMENT_TEMPLATES: dict[str, dict[str, Any]] = {
    "zju-grading-rubric-v1": {
        "schema_version": COMPANION_TEMPLATE_SCHEMA,
        "template_id": "zju-grading-rubric-v1",
        "template_version": 1,
        "document_type": "grading_rubric",
        "name": "评分细则",
        "name_en": "Grading rubric",
        "description": "按考核项目、比例与给分说明生成正式评分文件",
        "description_en": "Create a formal grading document from components, weights, and rules",
        "institution": "浙江大学",
        "audiences": ["school", "student"],
        "form_kind": "grading_rubric",
        "singleton": True,
        "export_formats": ["docx", "md"],
    },
    "zju-exam-course-material-checklist-v1": {
        "schema_version": COMPANION_TEMPLATE_SCHEMA,
        "template_id": "zju-exam-course-material-checklist-v1",
        "template_version": 1,
        "document_type": "course_material_checklist",
        "name": "考试课程材料自查清单",
        "name_en": "Exam-course material checklist",
        "description": "逐项核对考试课程材料是否齐全并生成归档清单",
        "description_en": "Check required exam-course records and create the archive checklist",
        "institution": "浙江大学",
        "audiences": ["school"],
        "form_kind": "material_checklist",
        "singleton": True,
        "export_formats": ["docx", "md"],
    },
}


class CompanionDocumentError(ValueError):
    """A template input or stored companion document is invalid."""


def list_templates(course: dict[str, Any]) -> list[dict[str, Any]]:
    """Return stable template contracts with course-aware default inputs."""
    return [
        {**deepcopy(template), "default_inputs": default_inputs(template_id, course)}
        for template_id, template in COMPANION_DOCUMENT_TEMPLATES.items()
    ]


def default_inputs(template_id: str, course: dict[str, Any]) -> dict[str, Any]:
    template = _template(template_id)
    profile = course.get("course_profile") if isinstance(course.get("course_profile"), dict) else {}
    request = course.get("generation_request") if isinstance(course.get("generation_request"), dict) else {}
    brief = request.get("teacher_course_brief") if isinstance(request.get("teacher_course_brief"), dict) else {}
    course_name = _text(course.get("course_name") or course.get("title") or "")
    course_code = _text(profile.get("course_code") or brief.get("course_code") or course.get("course_code"))
    academic_year = _text(course.get("academic_year") or "")
    term = _text(course.get("term") or "")
    if not academic_year or not term:
        academic_term = _text(brief.get("academic_term"))
        if academic_term:
            parts = academic_term.split(maxsplit=1)
            academic_year = academic_year or parts[0]
            term = term or (parts[1] if len(parts) > 1 else "")

    if template["form_kind"] == "grading_rubric":
        title_course_name = course_name if course_name.startswith("《") and course_name.endswith("》") else f"《{course_name}》"
        return {
            "title": f"{title_course_name}课程成绩评定细则" if course_name else "课程成绩评定细则",
            "course_name": course_name,
            "teacher_name": "",
            "effective_date": "",
            "components": deepcopy(GRADING_COMPONENTS),
            "special_rules": "作业延迟或补交（需合理理由）原则上按80%折算成绩。MOOC超时未交记零分；补交或补考（需合理理由）原则上按原得分50%计分。",
        }
    return {
        "title": "考试课程材料自查清单",
        "college_name": "计算机科学与技术学院",
        "course_name": course_name,
        "course_code": course_code,
        "academic_year": academic_year,
        "term": term,
        "exam_time": "",
        "course_type": _text(profile.get("course_category") or ""),
        "teacher_name": "",
        "submitted_at": "",
        "items": [
            {"item_id": item_id, "completed": False, "notes": ""}
            for item_id, _ in CHECKLIST_ITEMS
        ],
    }


def compile_document(
    template_id: str,
    inputs: dict[str, Any],
    course: dict[str, Any],
) -> dict[str, Any]:
    """Validate template inputs and compile a deterministic formal document."""
    template = _template(template_id)
    defaults = default_inputs(template_id, course)
    merged = {**defaults, **deepcopy(inputs or {})}
    if template["form_kind"] == "grading_rubric":
        normalized = _normalize_grading_inputs(merged)
        markdown = _grading_markdown(normalized)
    else:
        normalized = _normalize_checklist_inputs(merged)
        markdown = _checklist_markdown(normalized)
    return {
        "template": deepcopy(template),
        "inputs": normalized,
        "rendered_markdown": markdown,
    }


class CompanionDocumentRepository:
    """Persist one versioned course document for every singleton template."""

    def __init__(self, root_dir: str | Path | None = None) -> None:
        self.root_dir = Path(root_dir or Path(DATA_DIR) / "companion_documents")
        self.root_dir.mkdir(parents=True, exist_ok=True)

    def list_course(self, course_id: str) -> list[dict[str, Any]]:
        course_dir = self.root_dir / _storage_id(course_id)
        if not course_dir.exists():
            return []
        documents = [
            document
            for path in course_dir.iterdir()
            if path.is_dir()
            for document in [self._load_current(course_id, path.name)]
            if document is not None
        ]
        documents.sort(key=lambda item: str(item.get("updated_at") or ""), reverse=True)
        return deepcopy(documents)

    def load(self, course_id: str, document_id: str) -> dict[str, Any] | None:
        return deepcopy(self._load_current(course_id, document_id))

    def save_revision(
        self,
        *,
        course_id: str,
        actor_id: str,
        compiled: dict[str, Any],
    ) -> dict[str, Any]:
        normalized_course_id = _storage_id(course_id)
        template = compiled["template"]
        existing = next(
            (
                item
                for item in self.list_course(normalized_course_id)
                if item.get("template_id") == template["template_id"]
            ),
            None,
        )
        document_id = str(existing.get("document_id") or "") if existing else f"compdoc_{uuid4().hex}"
        now = _now()
        revision_number = int(existing.get("revision_number") or 0) + 1 if existing else 1
        signature_payload = {
            "template_id": template["template_id"],
            "template_version": template["template_version"],
            "inputs": compiled["inputs"],
            "rendered_markdown": compiled["rendered_markdown"],
        }
        revision_id = f"compdocrev_{_stable_digest(signature_payload)}"
        if existing and existing.get("revision_id") == revision_id:
            return existing
        document = {
            "schema_version": COMPANION_DOCUMENT_SCHEMA,
            "document_id": document_id,
            "course_id": normalized_course_id,
            "template_id": template["template_id"],
            "template_version": template["template_version"],
            "document_type": template["document_type"],
            "title": _text(compiled["inputs"].get("title") or template["name"]),
            "status": "ready",
            "revision_id": revision_id,
            "revision_number": revision_number,
            "inputs": deepcopy(compiled["inputs"]),
            "rendered_markdown": compiled["rendered_markdown"],
            "created_by": _text(actor_id)[:200],
            "created_at": _text(existing.get("created_at")) if existing else now,
            "updated_at": now,
        }
        self._write_revision(document)
        return deepcopy(document)

    def _load_current(self, course_id: str, document_id: str) -> dict[str, Any] | None:
        normalized_course_id = _storage_id(course_id)
        normalized_document_id = _storage_id(document_id)
        pointer = self.root_dir / normalized_course_id / normalized_document_id / "current.json"
        if not pointer.exists():
            return None
        revision_id = _text(_read_json(pointer).get("revision_id"))
        if not revision_id:
            return None
        revision_path = pointer.parent / "revisions" / f"{_storage_id(revision_id)}.json"
        if not revision_path.exists():
            return None
        document = _read_json(revision_path)
        if document.get("course_id") != normalized_course_id:
            raise CompanionDocumentError("配套文档课程范围不合法")
        return document

    def _write_revision(self, document: dict[str, Any]) -> None:
        course_id = _storage_id(_text(document.get("course_id")))
        document_id = _storage_id(_text(document.get("document_id")))
        revision_id = _storage_id(_text(document.get("revision_id")))
        directory = self.root_dir / course_id / document_id
        revision_path = directory / "revisions" / f"{revision_id}.json"
        if not revision_path.exists():
            _atomic_write(revision_path, document)
        _atomic_write(directory / "current.json", {"revision_id": revision_id})


def export_document(document: dict[str, Any], export_format: str) -> tuple[bytes, str, str]:
    normalized_format = _text(export_format).lower()
    if normalized_format == "md":
        return (
            _text(document.get("rendered_markdown")).encode("utf-8"),
            "text/markdown; charset=utf-8",
            f"{_safe_filename(document.get('title'))}.md",
        )
    if normalized_format != "docx":
        raise CompanionDocumentError("暂不支持该导出格式")
    if document.get("document_type") == "grading_rubric":
        payload = _grading_docx(document.get("inputs") or {})
    elif document.get("document_type") == "course_material_checklist":
        payload = _checklist_docx(document.get("inputs") or {})
    else:
        raise CompanionDocumentError("配套文档类型不支持导出")
    return payload, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{_safe_filename(document.get('title'))}.docx"


def _normalize_grading_inputs(values: dict[str, Any]) -> dict[str, Any]:
    components = values.get("components")
    if not isinstance(components, list) or not 1 <= len(components) <= 20:
        raise CompanionDocumentError("评分细则至少需要一个考核项目")
    normalized_components = []
    for index, raw in enumerate(components, start=1):
        if not isinstance(raw, dict):
            raise CompanionDocumentError("考核项目格式不合法")
        name = _text(raw.get("name")).strip()
        if not name:
            raise CompanionDocumentError(f"第{index}个考核项目缺少名称")
        try:
            weight = round(float(raw.get("weight") or 0), 2)
        except (TypeError, ValueError) as exc:
            raise CompanionDocumentError(f"第{index}个考核项目比例不合法") from exc
        if weight <= 0 or weight > 100:
            raise CompanionDocumentError(f"第{index}个考核项目比例必须大于0且不超过100")
        normalized_components.append({
            "component_id": _text(raw.get("component_id") or f"component_{index}"),
            "name": name[:160],
            "weight": weight,
            "scope": _text(raw.get("scope"))[:80],
            "details": _text(raw.get("details"))[:4000],
        })
    total = round(sum(float(item["weight"]) for item in normalized_components), 2)
    if abs(total - 100) > 0.001:
        raise CompanionDocumentError(f"各考核项目比例合计必须为100%，当前为{total:g}%")
    return {
        "title": _required_text(values.get("title"), "文档标题", 240),
        "course_name": _required_text(values.get("course_name"), "课程名称", 240),
        "teacher_name": _text(values.get("teacher_name"))[:120],
        "effective_date": _text(values.get("effective_date"))[:80],
        "components": normalized_components,
        "special_rules": _text(values.get("special_rules"))[:4000],
    }


def _normalize_checklist_inputs(values: dict[str, Any]) -> dict[str, Any]:
    raw_items = values.get("items") if isinstance(values.get("items"), list) else []
    states = {
        _text(item.get("item_id")): item
        for item in raw_items
        if isinstance(item, dict) and _text(item.get("item_id"))
    }
    return {
        "title": _required_text(values.get("title"), "文档标题", 240),
        "college_name": _text(values.get("college_name"))[:240],
        "course_name": _required_text(values.get("course_name"), "课程名称", 240),
        "course_code": _text(values.get("course_code"))[:80],
        "academic_year": _text(values.get("academic_year"))[:80],
        "term": _text(values.get("term"))[:80],
        "exam_time": _text(values.get("exam_time"))[:120],
        "course_type": _text(values.get("course_type"))[:120],
        "teacher_name": _text(values.get("teacher_name"))[:120],
        "submitted_at": _text(values.get("submitted_at"))[:80],
        "items": [
            {
                "item_id": item_id,
                "label": label,
                "completed": bool(states.get(item_id, {}).get("completed")),
                "notes": _text(states.get(item_id, {}).get("notes"))[:500],
            }
            for item_id, label in CHECKLIST_ITEMS
        ],
    }


def _grading_markdown(values: dict[str, Any]) -> str:
    lines = [f"# {values['title']}", "", f"课程名称：{values['course_name']}", "", "本课程成绩包括以下部分：", ""]
    for index, item in enumerate(values["components"], start=1):
        weight = f"{float(item['weight']):g}"
        scope = f"，{item['scope']}" if item.get("scope") else ""
        lines.extend([f"## {index}. {item['name']}（{weight}%{scope}）", "", _text(item.get("details")), ""])
    if values.get("special_rules"):
        lines.extend(["## 补交与特殊情况", "", values["special_rules"], ""])
    lines.extend([f"任课教师：{values.get('teacher_name') or '____________'}", "", f"日期：{values.get('effective_date') or '____________'}"])
    return "\n".join(lines).strip() + "\n"


def _checklist_markdown(values: dict[str, Any]) -> str:
    prefix = f"{values.get('college_name')}" if values.get("college_name") else ""
    lines = [f"# {prefix}{values['title']}", "", f"课程名称：{values['course_name']}　课程代码：{values.get('course_code') or '________'}", f"开课时间：{values.get('academic_year') or '________'}学年　{values.get('term') or '____'}学期　考试时间：{values.get('exam_time') or '________'}", f"课程类型：{values.get('course_type') or '________'}", "", "## 考试材料清单（按教学班整理）", ""]
    for index, item in enumerate(values["items"], start=1):
        marker = "x" if item["completed"] else " "
        note = f"（{item['notes']}）" if item.get("notes") else ""
        lines.append(f"- [{marker}] {index}. {item['label']}{note}")
    lines.extend(["", "注：本表逐项打勾确认，并随同其他资料一并上交至学院本科生科存档。", "", f"课程教师签名：{values.get('teacher_name') or '____________'}　递交日期：{values.get('submitted_at') or '____________'}"])
    return "\n".join(lines).strip() + "\n"


def _grading_docx(values: dict[str, Any]) -> bytes:
    document = _base_document()
    _add_title(document, _text(values.get("title")))
    paragraph = document.add_paragraph(f"课程名称：{_text(values.get('course_name'))}")
    _style_paragraph(paragraph)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "考核项目", "比例", "对象", "评分项目与提交物", "备注"]
    for index, label in enumerate(headers):
        _set_cell(table.cell(0, index), label, bold=True, centered=True)
    for index, item in enumerate(values.get("components") or [], start=1):
        row = table.add_row().cells
        _set_cell(row[0], str(index), centered=True)
        _set_cell(row[1], _text(item.get("name")))
        _set_cell(row[2], f"{float(item.get('weight') or 0):g}%", centered=True)
        _set_cell(row[3], _text(item.get("scope")), centered=True)
        _set_cell(row[4], _text(item.get("details")))
        _set_cell(row[5], "")
    total_row = table.add_row().cells
    _set_cell(total_row[0].merge(total_row[1]), "合计", bold=True, centered=True)
    _set_cell(total_row[2], "100%", bold=True, centered=True)
    _set_cell(total_row[3].merge(total_row[5]), "", centered=True)
    if values.get("special_rules"):
        paragraph = document.add_paragraph(f"说明：{values['special_rules']}")
        _style_paragraph(paragraph)
    sign = document.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run(f"任课教师：{values.get('teacher_name') or '____________'}\n日期：{values.get('effective_date') or '____________'}")
    _style_paragraph(sign)
    return _document_bytes(document)


def _checklist_docx(values: dict[str, Any]) -> bytes:
    document = _base_document()
    _add_title(document, f"{values.get('college_name') or ''}{values.get('title') or ''}")
    metadata = document.add_table(rows=3, cols=4)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = [
        (0, 0, "课程名称", True), (0, 1, _text(values.get("course_name")), False),
        (0, 2, "课程代码", True), (0, 3, _text(values.get("course_code")), False),
        (1, 0, "开课时间", True), (1, 1, f"{values.get('academic_year') or ''} {values.get('term') or ''}", False),
        (1, 2, "考试时间", True), (1, 3, _text(values.get("exam_time")), False),
        (2, 0, "课程类型", True), (2, 1, _text(values.get("course_type")), False),
        (2, 2, "任课教师", True), (2, 3, _text(values.get("teacher_name")), False),
    ]
    for row, column, value, bold in cells:
        _set_cell(metadata.cell(row, column), value, bold=bold, centered=bold)
    paragraph = document.add_paragraph("考试材料清单（按教学班整理）：")
    _style_paragraph(paragraph, bold=True)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(["序号", "材料名称", "是否完成", "备注"]):
        _set_cell(table.cell(0, index), label, bold=True, centered=True)
    for index, item in enumerate(values.get("items") or [], start=1):
        row = table.add_row().cells
        _set_cell(row[0], str(index), centered=True)
        _set_cell(row[1], _text(item.get("label")))
        _set_cell(row[2], "√" if item.get("completed") else "□", centered=True)
        _set_cell(row[3], _text(item.get("notes")))
    note = document.add_paragraph("注：本表逐项打勾确认，并随同其他资料一并上交至学院本科生科存档。")
    _style_paragraph(note)
    sign = document.add_paragraph()
    sign.add_run(f"课程教师签名：{values.get('teacher_name') or '____________'}　　递交日期：{values.get('submitted_at') or '____________'}")
    _style_paragraph(sign)
    return _document_bytes(document)


def _base_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    return document


def _add_title(document: Document, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(value)
    run.bold = True
    _style_run(run, 16)


def _style_paragraph(paragraph: Any, *, bold: bool = False) -> None:
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.35
    for run in paragraph.runs:
        run.bold = run.bold or bold
        _style_run(run, 10.5)


def _style_run(run: Any, size: float) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def _set_cell(cell: Any, value: str, *, bold: bool = False, centered: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(value)
    run.bold = bold
    _style_run(run, 9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _document_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _template(template_id: str) -> dict[str, Any]:
    template = COMPANION_DOCUMENT_TEMPLATES.get(_text(template_id))
    if template is None:
        raise CompanionDocumentError("配套文档模板不存在")
    return template


def _required_text(value: Any, label: str, maximum: int) -> str:
    normalized = _text(value).strip()
    if not normalized:
        raise CompanionDocumentError(f"{label}不能为空")
    return normalized[:maximum]


def _text(value: Any) -> str:
    return "" if value is None else str(value)


def _storage_id(value: str) -> str:
    normalized = _text(value).strip()
    if not _ID_RE.fullmatch(normalized):
        raise CompanionDocumentError("配套文档标识不合法")
    return normalized


def _stable_digest(value: dict[str, Any]) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:32]


def _read_json(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as handle:
        value = json.load(handle)
    if not isinstance(value, dict):
        raise CompanionDocumentError("配套文档仓库数据格式不合法")
    return value


def _atomic_write(path: Path, value: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8") as handle:
        json.dump(value, handle, ensure_ascii=False, indent=2)
        handle.flush()
    temporary.replace(path)


def _safe_filename(value: Any) -> str:
    normalized = re.sub(r"[\\/:*?\"<>|]+", "-", _text(value).strip()).strip(". ")
    return normalized[:120] or "配套文档"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


companion_document_repository = CompanionDocumentRepository()


__all__ = [
    "COMPANION_DOCUMENT_SCHEMA",
    "COMPANION_DOCUMENT_TEMPLATES",
    "CompanionDocumentError",
    "CompanionDocumentRepository",
    "companion_document_repository",
    "compile_document",
    "export_document",
    "list_templates",
]
