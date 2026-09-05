import io
import math
import re
from typing import Any

import docx
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from wordcloud import WordCloud

from common.utils.logger import get_logger

logger = get_logger(__name__)

# ── Matplotlib 中文字体配置 ──
matplotlib.rcParams["font.family"] = ["Arial Unicode MS", "PingFang HK", "Heiti TC", "sans-serif"]
matplotlib.rcParams["axes.unicode_minus"] = False

_XML_BAD_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]")

COLOR_HEADER = RGBColor(0x33, 0x33, 0x33)
COLOR_PASS = RGBColor(0x67, 0xC2, 0x3A)
COLOR_WARN = RGBColor(0xFA, 0xAD, 0x14)
COLOR_FAIL = RGBColor(0xF5, 0x6C, 0x6C)
COLOR_GRAY = RGBColor(0x90, 0x93, 0x99)


def _clean_text(text: Any) -> str:
    if text is None:
        return ""
    return _XML_BAD_CHARS.sub("", str(text))


def _score_color(score: float) -> RGBColor:
    if score >= 80:
        return COLOR_PASS
    if score >= 60:
        return COLOR_WARN
    return COLOR_FAIL


def _set_run_font(run, font_name: str = "微软雅黑", size: int | None = None, color: RGBColor | None = None, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold


def _add_styled_paragraph(doc: docx.Document, text: str, size: int = 10, bold: bool = False, color: RGBColor | None = None, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(_clean_text(text))
    _set_run_font(run, size=size, bold=bold, color=color)
    return p


def _add_heading_styled(doc: docx.Document, text: str, level: int = 1):
    heading = doc.add_heading(_clean_text(text), level=level)
    for run in heading.runs:
        _set_run_font(run, size=16 if level == 1 else 13, bold=True, color=COLOR_HEADER)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    return heading


# ── 图表生成 ──

def _generate_radar_chart_image(radar_data: dict[str, float]) -> bytes | None:
    """生成六维雷达图 PNG 图片字节。"""
    if not radar_data:
        return None

    labels = list(radar_data.keys())
    values = list(radar_data.values())
    num_vars = len(labels)

    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    values += values[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor("white")

    # 背景网格
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_thetagrids(np.degrees(angles[:-1]), labels, fontsize=12)
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(["20", "40", "60", "80", "100"], color="gray", size=8)
    ax.grid(color="#E0E0E0", linestyle="-", linewidth=0.8)

    # 数据区域
    color_fill = "#4A90D9"
    color_line = "#2E6AAD"
    ax.fill(angles, values, color=color_fill, alpha=0.25)
    ax.plot(angles, values, color=color_line, linewidth=2.5, marker="o", markersize=6)

    # 顶点标注分数
    for angle, value, label in zip(angles[:-1], values[:-1], labels):
        ax.annotate(f"{value}", xy=(angle, value), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=10, fontweight="bold", color=color_line)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _generate_wordcloud_image(knowledge_graph: dict[str, Any]) -> bytes | None:
    """从知识点图谱生成词云 PNG 图片字节。"""
    nodes = knowledge_graph.get("nodes") if isinstance(knowledge_graph, dict) else []
    if not isinstance(nodes, list) or not nodes:
        return None

    # 收集词频（名称 + 类型权重）
    frequencies = {}
    for node in nodes:
        if not isinstance(node, dict):
            continue
        name = node.get("name", "")
        if name:
            # 优先使用显式词频权重（V2 词云 {word, weight}）；缺省时按类型粗略赋权（超星原始结构）
            weight = node.get("weight")
            if not isinstance(weight, (int, float)) or weight <= 0:
                weight = 2 if node.get("type") == "core" else 1
            frequencies[name] = frequencies.get(name, 0) + weight

    if not frequencies:
        return None

    wordcloud = WordCloud(
        width=800, height=400,
        background_color="white",
        font_path=None,  # 使用系统默认字体
        colormap="Blues",
        max_words=100,
        relative_scaling=0.5,
        min_font_size=10,
        max_font_size=120,
    ).generate_from_frequencies(frequencies)

    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor("white")
    ax.imshow(wordcloud, interpolation="bilinear")
    ax.axis("off")
    ax.set_title("知识点词云", fontsize=14, fontweight="bold", pad=10, color="#333333")

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white", pad_inches=0.1)
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _generate_volume_chart_image(db_result: dict[str, Any]) -> bytes | None:
    """生成音量趋势折线图 PNG 图片字节。"""
    if not isinstance(db_result, dict):
        return None

    data = db_result.get("data", {})
    if not isinstance(data, dict):
        return None

    values = data.get("result")
    if not isinstance(values, list) or len(values) < 2:
        return None

    # 采样点过多时降采样
    max_points = 200
    if len(values) > max_points:
        step = math.ceil(len(values) / max_points)
        values = values[::step]

    x = list(range(len(values)))

    fig, ax = plt.subplots(figsize=(8, 3.5))
    fig.patch.set_facecolor("white")

    ax.fill_between(x, values, alpha=0.3, color="#4A90D9")
    ax.plot(x, values, color="#2E6AAD", linewidth=1.5)
    ax.set_xlabel("时间", fontsize=10, color="#666666")
    ax.set_ylabel("声压级 (dB)", fontsize=10, color="#666666")
    ax.set_title("音量变化趋势", fontsize=14, fontweight="bold", pad=10, color="#333333")
    ax.grid(True, linestyle="--", alpha=0.4)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # 标注 avg / max / min
    avg_spl = data.get("avg_spl")
    if isinstance(avg_spl, (int, float)):
        ax.axhline(y=avg_spl, color="#F5A623", linestyle="--", linewidth=1, label=f"平均: {avg_spl:.1f} dB")
        ax.legend(loc="upper right", fontsize=9)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _generate_wh_pie(cats: dict[str, int]) -> bytes | None:
    """生成五何互动占比饼图 PNG 图片字节。"""
    if not cats:
        return None

    labels = list(cats.keys())
    sizes = list(cats.values())
    total = sum(sizes)

    # 蓝色系配色
    colors = ["#4A90D9", "#5BA3E8", "#7DBCF2", "#A8D4F7", "#D0E8FA"]

    fig, ax = plt.subplots(figsize=(5, 5))
    fig.patch.set_facecolor("white")

    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        autopct=lambda pct: f"{pct:.1f}%\n({int(pct/100.*total)})",
        startangle=90,
        colors=colors[:len(labels)],
        textprops={"fontsize": 10},
        pctdistance=0.65,
    )
    for autotext in autotexts:
        autotext.set_fontsize(9)
        autotext.set_color("#333333")

    ax.set_title("五何互动占比", fontsize=14, fontweight="bold", pad=10, color="#333333")
    ax.axis("equal")

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _generate_teaching_segments_bar(teach_summary: list[dict]) -> bytes | None:
    """生成教学环节时长分布条形图。"""
    if not teach_summary:
        return None

    segments = []
    for item in teach_summary:
        if not isinstance(item, dict):
            continue
        for seg in item.get("file_structure", []):
            if isinstance(seg, dict):
                segments.append(seg)

    if not segments:
        return None

    # 按类型聚合时长
    from collections import defaultdict
    type_durations = defaultdict(float)
    for seg in segments:
        t = seg.get("type", "未知")
        # 解析时长（简单处理：end_time - start_time 的秒数）
        try:
            start = seg.get("start_time", "0:00")
            end = seg.get("end_time", "0:00")
            # 简单解析 MM:SS 或 HH:MM:SS
            def _to_sec(ts: str) -> float:
                parts = ts.split(":")
                if len(parts) == 2:
                    return int(parts[0]) * 60 + float(parts[1])
                elif len(parts) == 3:
                    return int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2])
                return 0
            dur = _to_sec(end) - _to_sec(start)
            if dur > 0:
                type_durations[t] += dur
        except Exception:
            pass

    if not type_durations:
        return None

    labels = list(type_durations.keys())
    durations = [type_durations[k] / 60 for k in labels]  # 转分钟

    fig, ax = plt.subplots(figsize=(7, 4))
    fig.patch.set_facecolor("white")

    colors = plt.cm.Blues(np.linspace(0.4, 0.8, len(labels)))
    bars = ax.barh(labels, durations, color=colors, height=0.6)
    ax.set_xlabel("时长 (分钟)", fontsize=10, color="#666666")
    ax.set_title("教学环节时长分布", fontsize=14, fontweight="bold", pad=10, color="#333333")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    for bar, val in zip(bars, durations):
        ax.text(val + 0.3, bar.get_y() + bar.get_height() / 2, f"{val:.1f}",
                va="center", fontsize=9, color="#333333")

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ── AI 建议生成 ──

async def _generate_ai_suggestions(analysis_result: dict[str, Any]) -> str:
    """调用 LLM 基于完整分析数据给出教学改进建议。"""
    radar = analysis_result.get("radar_chart") or {}
    teach_wh = analysis_result.get("teach_wh") or []
    db_result = analysis_result.get("teach_db_result") or {}
    teach_summary = analysis_result.get("teach_summary") or []
    class_summary = analysis_result.get("class_summary") or {}
    class_edu = analysis_result.get("class_education_summary") or {}

    # 1. 雷达图
    radar_text = "\n".join(f"- {k}: {v}分" for k, v in radar.items()) if radar else "暂无"

    # 2. 五何互动
    wh_stats = {}
    if isinstance(teach_wh, list):
        for item in teach_wh:
            if isinstance(item, dict):
                cat = item.get("category", "未知")
                wh_stats[cat] = wh_stats.get(cat, 0) + 1
    wh_text = "\n".join(f"- {k}: {v}次" for k, v in wh_stats.items()) if wh_stats else "暂无"

    # 3. 音量语速
    db_data = db_result.get("data", {}) if isinstance(db_result, dict) else {}
    db_text = ""
    if isinstance(db_data, dict):
        db_text = f"平均声压级: {db_data.get('avg_spl', '未知')} dB, 语速: {db_data.get('speed', '未知')} 字/分钟"

    # 4. 教学环节类型
    seg_types = []
    if isinstance(teach_summary, list):
        for item in teach_summary:
            if isinstance(item, dict):
                for seg in item.get("file_structure", []):
                    if isinstance(seg, dict):
                        t = seg.get("type", "")
                        if t:
                            seg_types.append(t)
    from collections import Counter
    seg_counter = Counter(seg_types)
    seg_text = "\n".join(f"- {k}: {v}段" for k, v in seg_counter.most_common()) if seg_counter else "暂无"

    # 5. 课堂总结
    summary_text = ""
    if isinstance(class_summary, dict):
        summary_text = class_summary.get("summary", "") or "暂无"

    # 6. 思政融合
    edu_count = 0
    if isinstance(class_edu, dict):
        edu_list = class_edu.get("summary", [])
        if isinstance(edu_list, list):
            edu_count = len(edu_list)

    prompt = f"""你是一位资深教学督导专家。请综合以下完整的课堂教学分析数据，给出3-5条精准、具体、可落地的教学改进建议。

要求：
- 每条建议控制在40字以内
- 使用教师视角（"你"）
- 针对数据中的明显短板提出改进方向
- 不要空泛套话

【六维雷达图评分】（满分100）
{radar_text}

【教学环节类型分布】
{seg_text}

【五何互动统计】
{wh_text}

【音量与语速】
{db_text}

【课堂总结摘要】
{summary_text}

【思政融合】
- 思政元素出现次数: {edu_count}

请直接输出建议列表，每条以"1. "、"2. "开头。不要添加总结性段落。"""

    try:
        from agents.llm import complete_chat
        suggestion = await complete_chat(
            system_prompt="你是一位资深教学督导专家，擅长基于多维度课堂数据分析给出精准的教学改进建议。",
            user_prompt=prompt,
        )
        return suggestion
    except Exception:
        logger.exception("AI 建议生成失败")
        return "AI 建议生成暂时不可用，请稍后重试。"


# ── Word 文档生成 ──

def _add_image(doc: docx.Document, image_bytes: bytes | None, width: Inches | None = None):
    """向文档插入图片。"""
    if not image_bytes:
        return
    try:
        doc.add_picture(io.BytesIO(image_bytes), width=width or Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        logger.warning("插入图片失败")


def _is_v2_analysis_schema(result: dict[str, Any]) -> bool:
    """判断 analysis_result 是否为新版 V2 五维结构（与前端 isNewVideoAnalysisFormat 口径一致）。"""
    if not isinstance(result, dict):
        return False
    if isinstance(result.get("radar_data"), list):
        return True
    return any(
        key in result
        for key in (
            "teaching_expression",
            "teaching_design",
            "knowledge_presentation",
            "interaction_quality",
            "ideological_integration",
        )
    )


def _map_v2_to_export_schema(result: dict[str, Any]) -> dict[str, Any]:
    """将线上入库的 V2 五维结构映射为本导出器既有渲染逻辑所需的字段。

    线上 analysis_result 实际是 V2 结构（teaching_expression / teaching_design /
    knowledge_presentation / interaction_quality / ideological_integration /
    radar_data / ai_summary，见 analysis/orchestrator.py 与 local_analysis/adapter.py），
    而导出器原按超星原始字段读取（radar_chart / teach_db_result.data /
    knowledge_graph.nodes / teach_summary[].file_structure / teach_wh[].category），
    二者不匹配导致导出内容几乎为空。此处口径与前端 normalizeNewVideoAnalysisResult 对齐，
    使导出报告与报告页内容一致。
    """
    expression = result.get("teaching_expression") or {}
    volume = expression.get("volume_analysis") or {}
    speech = expression.get("speech_rate_analysis") or {}
    design = result.get("teaching_design") or {}
    knowledge = result.get("knowledge_presentation") or {}
    interaction = result.get("interaction_quality") or {}
    ai_summary = result.get("ai_summary") or {}

    # 雷达图：[{dimension, score}] → {维度: 分数}；剔除「综合得分」（导出器自行求均值展示）
    radar_chart: dict[str, float] = {}
    for item in result.get("radar_data") or []:
        if not isinstance(item, dict):
            continue
        dim = str(item.get("dimension") or item.get("label") or item.get("name") or "").strip()
        score = item.get("score", item.get("value"))
        if dim and dim != "综合得分" and isinstance(score, (int, float)):
            radar_chart[dim] = score

    # 教学改进建议：综合评价置顶 + 建议列表
    suggestions = [
        s.strip()
        for s in [ai_summary.get("summary"), *(ai_summary.get("suggestions") or [])]
        if isinstance(s, str) and s.strip()
    ]

    # 教学环节：teaching_design.segments → teach_summary[].file_structure
    file_structure = [
        {
            "type": seg.get("type", ""),
            "content": seg.get("content", ""),
            "start_time": seg.get("start_time", ""),
            "end_time": seg.get("end_time", ""),
            "keypoint": seg.get("keypoint_name", ""),
        }
        for seg in (design.get("segments") or [])
        if isinstance(seg, dict)
    ]
    teach_summary = [{"file_structure": file_structure}] if file_structure else []

    # 知识点词云：word_cloud[{word, weight}] → knowledge_graph.nodes[{name, weight}]
    nodes = [
        {"name": w.get("word", ""), "weight": w.get("weight", 1)}
        for w in (knowledge.get("word_cloud") or [])
        if isinstance(w, dict) and w.get("word")
    ]
    knowledge_graph = {"nodes": nodes}

    # 音量趋势 + 语速：teaching_expression → teach_db_result.data
    teach_db_result = {
        "data": {
            "result": volume.get("result") or [],
            "avg_spl": volume.get("avg_spl"),
            "speed": speech.get("avg_cpm"),
        }
    }

    # 五何互动：wh_distribution{类别:{count}} → teach_wh[{category}]（按 count 展开以便计数/占比）
    teach_wh: list[dict[str, Any]] = []
    wh_dist = interaction.get("wh_distribution") or {}
    if isinstance(wh_dist, dict):
        for cat, info in wh_dist.items():
            count = 0
            if isinstance(info, dict):
                try:
                    count = int(info.get("count", 0))
                except (TypeError, ValueError):
                    count = 0
            teach_wh.extend({"category": cat} for _ in range(max(0, count)))

    return {
        "audio_duration": volume.get("total_duration") or result.get("audio_duration"),
        "radar_chart": radar_chart,
        "suggestions": suggestions,
        "teach_summary": teach_summary,
        "knowledge_graph": knowledge_graph,
        "teach_db_result": teach_db_result,
        "teach_wh": teach_wh,
    }


async def generate_video_analysis_report(video_name: str, analysis_result: dict[str, Any]) -> bytes:
    """
    将视频分析结果生成图文报告 Word 文件。
    包含雷达图、词云、音量图、AI 建议，尽量减少纯文字表格。
    """
    # 线上入库为 V2 五维结构，先映射为导出器所需的字段口径（旧版超星结构则原样透传）
    if _is_v2_analysis_schema(analysis_result):
        analysis_result = _map_v2_to_export_schema(analysis_result)

    doc = docx.Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(10)

    # ── 封面 ──
    title = doc.add_heading("课堂教学分析报告", level=0)
    for run in title.runs:
        _set_run_font(run, size=22, bold=True, color=COLOR_HEADER)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)

    _add_styled_paragraph(doc, f"视频名称：{video_name}", size=12, color=COLOR_GRAY, space_after=24)

    # 时长信息
    audio_duration = analysis_result.get("audio_duration")
    if audio_duration:
        m = int(audio_duration // 60)
        s = int(audio_duration % 60)
        duration_str = f"{m}分{s}秒" if m < 60 else f"{m // 60}小时{m % 60}分{s}秒"
        _add_styled_paragraph(doc, f"视频时长：{duration_str}", size=11, color=COLOR_GRAY, space_after=12)

    # ── 1. 综合评分 + 雷达图 ──
    radar_chart = analysis_result.get("radar_chart") or {}
    if radar_chart:
        _add_heading_styled(doc, "一、综合评分", level=1)

        # 雷达图
        radar_img = _generate_radar_chart_image(radar_chart)
        _add_image(doc, radar_img, width=Inches(5))

        # 平均分
        avg_score = sum(radar_chart.values()) / len(radar_chart)
        avg_p = doc.add_paragraph()
        avg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        avg_p.paragraph_format.space_before = Pt(8)
        avg_p.paragraph_format.space_after = Pt(12)
        avg_label = avg_p.add_run("综合评分：")
        _set_run_font(avg_label, size=12, bold=True)
        avg_value = avg_p.add_run(f"{avg_score:.1f}")
        _set_run_font(avg_value, size=16, bold=True, color=_score_color(avg_score))

    # ── 2. AI 教学建议 ──
    _add_heading_styled(doc, "二、教学改进建议", level=1)
    suggestions = analysis_result.get("suggestions") or []
    if suggestions:
        for line in suggestions:
            line = line.strip()
            if not line:
                continue
            _add_styled_paragraph(doc, line, size=10, space_after=6)
    else:
        _add_styled_paragraph(doc, "暂无建议数据", size=10, color=COLOR_GRAY, space_after=6)

    # ── 3. 教学环节分布 ──
    teach_summary = analysis_result.get("teach_summary") or []
    if teach_summary:
        _add_heading_styled(doc, "三、教学环节分布", level=1)
        seg_img = _generate_teaching_segments_bar(teach_summary)
        _add_image(doc, seg_img, width=Inches(5.5))

    # ── 4. 知识点词云 ──
    knowledge_graph = analysis_result.get("knowledge_graph") or {}
    if isinstance(knowledge_graph, dict) and knowledge_graph.get("nodes"):
        _add_heading_styled(doc, "四、知识点分布", level=1)
        cloud_img = _generate_wordcloud_image(knowledge_graph)
        _add_image(doc, cloud_img, width=Inches(5.5))

    # ── 5. 音量趋势 ──
    teach_db_result = analysis_result.get("teach_db_result") or {}
    if isinstance(teach_db_result, dict):
        volume_img = _generate_volume_chart_image(teach_db_result)
        if volume_img:
            _add_heading_styled(doc, "五、音量变化趋势", level=1)
            _add_image(doc, volume_img, width=Inches(5.5))

            # 简要指标
            db_data = teach_db_result.get("data", {})
            if isinstance(db_data, dict):
                metrics = []
                if db_data.get("avg_spl") is not None:
                    metrics.append(f"平均声压级：{db_data['avg_spl']} dB")
                if db_data.get("speed") is not None:
                    metrics.append(f"语速：{db_data['speed']} 字/分钟")
                if metrics:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    run = p.add_run("  |  ".join(metrics))
                    _set_run_font(run, size=9, color=COLOR_GRAY)

    # ── 6. 五何互动统计（表格 + 饼图）──
    teach_wh = analysis_result.get("teach_wh") or []
    if teach_wh:
        _add_heading_styled(doc, "六、五何互动统计", level=1)
        from collections import Counter
        cats = Counter(item.get("category", "未知") for item in teach_wh if isinstance(item, dict))
        if cats:
            total = sum(cats.values())
            wh_rows = []
            for cat, count in cats.most_common():
                pct = count / total * 100
                wh_rows.append([cat, str(count), f"{pct:.1f}%"])

            table = doc.add_table(rows=1 + len(wh_rows), cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "互动类型", "次数", "占比"
            for i, (cat, cnt, pct) in enumerate(wh_rows, start=1):
                row = table.rows[i].cells
                row[0].text = cat
                row[1].text = cnt
                row[2].text = pct
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 饼图
            pie_img = _generate_wh_pie(cats)
            _add_image(doc, pie_img, width=Inches(4.5))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
